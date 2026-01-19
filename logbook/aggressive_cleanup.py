from docx import Document
import os

def aggressive_header_cleanup():
    """Aggressively remove ALL duplicate headers"""
    print("="*70)
    print("AGGRESSIVE HEADER CLEANUP")
    print("="*70)
    
    filename = "Logbook Lengkap.docx"
    
    if not os.path.exists(filename):
        print(f"ERROR: {filename} not found!")
        return
    
    print(f"\n[1/4] Loading {filename}...")
    doc = Document(filename)
    
    print(f"\n[2/4] Finding first table position...")
    
    # Find position of first table in body
    first_table_pos = None
    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('tbl'):
            first_table_pos = i
            print(f"   First table at position {i}")
            break
    
    if first_table_pos is None:
        print("   ERROR: No tables found!")
        return
    
    # Keywords that indicate header content
    header_indicators = [
        "Nama", "NIM", "Program Studi", "Nomor HP", 
        "Dosen Pembimbing", "Lokasi Pelaksanaan", 
        "Waktu Pelaksanaan", "PROJECT INDEPENDENT",
        "Muhamad Farhan", "221420075", "Teknik Informatika",
        "083826383761", "Rasmila", "Warung Sate Madura",
        "06 Oktober 2025"
    ]
    
    print(f"\n[3/4] Removing ALL paragraphs with header content after first table...")
    
    elements_to_remove = []
    removed_samples = []
    
    # Iterate through all elements AFTER first table
    for i in range(first_table_pos + 1, len(doc.element.body)):
        element = doc.element.body[i]
        
        if element.tag.endswith('p'):
            # Get paragraph text
            para_index = None
            for j, p in enumerate(doc.paragraphs):
                if p._element == element:
                    para_index = j
                    break
            
            if para_index is not None:
                para = doc.paragraphs[para_index]
                text = para.text.strip()
                
                # Check if paragraph contains ANY header indicator
                if any(indicator in text for indicator in header_indicators):
                    elements_to_remove.append(element)
                    if len(removed_samples) < 10:
                        removed_samples.append(text[:70])
    
    print(f"   Found {len(elements_to_remove)} paragraphs to remove")
    print(f"\n   Sample removals:")
    for sample in removed_samples:
        print(f"      - {sample}")
    
    # Remove all marked elements
    for element in elements_to_remove:
        element.getparent().remove(element)
    
    print(f"\n[4/4] Saving cleaned file...")
    doc.save(filename)
    
    # Verify
    verify_doc = Document(filename)
    
    # Count remaining header-like paragraphs
    remaining = 0
    remaining_samples = []
    for para in verify_doc.paragraphs:
        if any(ind in para.text for ind in header_indicators):
            remaining += 1
            if len(remaining_samples) < 5:
                remaining_samples.append(para.text[:70])
    
    print(f"\n{'='*70}")
    print("✅ CLEANUP COMPLETE!")
    print(f"{'='*70}")
    print(f"  - Removed: {len(elements_to_remove)} paragraphs")
    print(f"  - Remaining header paragraphs: {remaining}")
    print(f"  - Total tables: {len(verify_doc.tables)}")
    
    if remaining > 0:
        print(f"\n   Remaining header paragraphs (should be at top only):")
        for sample in remaining_samples:
            print(f"      - {sample}")
    
    print("="*70)

if __name__ == "__main__":
    aggressive_header_cleanup()
