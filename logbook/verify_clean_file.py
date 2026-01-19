from docx import Document
import os

def verify_clean_file():
    """Verify the clean merged file"""
    print("="*70)
    print("VERIFYING CLEAN MERGED FILE")
    print("="*70)
    
    filename = "Logbook Lengkap - Clean.docx"
    
    if not os.path.exists(filename):
        print(f"ERROR: {filename} not found!")
        return
    
    doc = Document(filename)
    
    header_keywords = ["Nama", "NIM", "Program Studi", "Nomor HP", 
                      "Dosen Pembimbing", "Lokasi Pelaksanaan", "Waktu Pelaksanaan"]
    
    print(f"\n[1/3] Document Statistics:")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables: {len(doc.tables)}")
    
    # Count images
    image_count = 0
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
    except:
        pass
    print(f"   Total images: {image_count}")
    
    # Find ALL paragraphs with header keywords
    print(f"\n[2/3] Header Analysis:")
    header_locations = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(kw in text for kw in header_keywords):
            header_locations.append((i, text))
    
    print(f"   Total paragraphs with header keywords: {len(header_locations)}")
    
    # Show ALL header locations
    print(f"\n   All header keyword locations:")
    for idx, text in header_locations:
        print(f"      P{idx}: {text[:70]}")
    
    # Find first table position
    first_table_para = None
    para_count = 0
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            first_table_para = para_count
            break
        elif element.tag.endswith('p'):
            para_count += 1
    
    if first_table_para:
        headers_after_table = [h for h in header_locations if h[0] > first_table_para]
        
        print(f"\n   First table at paragraph ~{first_table_para}")
        print(f"   Headers BEFORE first table: {len(header_locations) - len(headers_after_table)}")
        print(f"   Headers AFTER first table: {len(headers_after_table)}")
        
        if headers_after_table:
            print(f"\n   ❌ PROBLEM: Found duplicate headers after first table!")
            for idx, text in headers_after_table[:5]:
                print(f"      P{idx}: {text[:70]}")
        else:
            print(f"\n   ✅ SUCCESS: No duplicate headers found!")
    
    # Table summary
    print(f"\n[3/3] Table Summary:")
    for i in range(min(3, len(doc.tables))):
        table = doc.tables[i]
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0
        print(f"   Table {i+1}: {rows} rows x {cols} cols")
    
    if len(doc.tables) > 3:
        print(f"   ... and {len(doc.tables) - 3} more tables")
    
    print(f"\n{'='*70}")
    print("VERIFICATION COMPLETE")
    print(f"{'='*70}")

if __name__ == "__main__":
    verify_clean_file()
