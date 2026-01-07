from docx import Document
import sys

def extract_logbook_structure(filename):
    doc = Document(filename)
    
    print(f"=== Analyzing {filename} ===\n")
    
    # Print paragraphs
    print("PARAGRAPHS (first 20):")
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"{i}: {para.text}")
    
    print("\n" + "="*50 + "\n")
    
    # Print tables
    print(f"TABLES: {len(doc.tables)} found\n")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n--- Table {table_idx + 1} ---")
        print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        
        # Print all rows
        for row_idx, row in enumerate(table.rows):
            cells_text = [cell.text.strip().replace('\n', ' ')[:100] for cell in row.cells]
            print(f"Row {row_idx}: {cells_text}")
        
        if table_idx >= 1:  # Only show first 2 tables in detail
            break

if __name__ == "__main__":
    extract_logbook_structure("logbook minggu 1.docx")
