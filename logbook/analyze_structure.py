from docx import Document
import os

def analyze_logbook_structure(filename):
    """Analyze the structure of a logbook file"""
    if not os.path.exists(filename):
        print(f"File not found: {filename}")
        return
    
    print(f"\n{'='*60}")
    print(f"Analyzing: {filename}")
    print('='*60)
    
    doc = Document(filename)
    
    # Count paragraphs before first table
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Show first 20 paragraphs (header section)
    print("\n--- Header Section (First 20 paragraphs) ---")
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            print(f"P{i}: {para.text[:80]}")
    
    # Show table structure
    print("\n--- Table Structure ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        print(f"  Rows: {len(table.rows)}")
        print(f"  Columns: {len(table.rows[0].cells) if table.rows else 0}")
        
        # Show first row (header)
        if table.rows:
            first_row = table.rows[0]
            headers = [cell.text.strip() for cell in first_row.cells]
            print(f"  Headers: {headers}")
    
    # Check for images
    print("\n--- Images in Document ---")
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            print(f"  Image found: {rel.target_ref}")
    print(f"Total images: {image_count}")

if __name__ == "__main__":
    # Analyze first and last week to understand structure
    analyze_logbook_structure("logbook minggu 1.docx")
    analyze_logbook_structure("logbook minggu 16.docx")
