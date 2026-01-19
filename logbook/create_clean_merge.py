from docx import Document
from copy import deepcopy
import os

def create_clean_merged_logbook():
    """Create clean merged logbook - header once, then all tables"""
    print("="*70)
    print("CREATING CLEAN MERGED LOGBOOK (New File)")
    print("="*70)
    
    # Create brand new document
    merged_doc = Document()
    
    # Step 1: Copy header from first file ONLY
    print("\n[1/3] Extracting header from logbook minggu 1...")
    first_file = "logbook minggu 1.docx"
    
    if not os.path.exists(first_file):
        print(f"ERROR: {first_file} not found!")
        return
    
    doc1 = Document(first_file)
    
    # Find where first table starts
    first_table_index = None
    for i, element in enumerate(doc1.element.body):
        if element.tag.endswith('tbl'):
            first_table_index = i
            break
    
    # Copy all elements BEFORE first table (this is the header)
    header_count = 0
    for i, element in enumerate(doc1.element.body):
        if i >= first_table_index:
            break
        
        # Deep copy the element
        new_element = deepcopy(element)
        merged_doc.element.body.append(new_element)
        header_count += 1
    
    print(f"   ✓ Copied {header_count} header elements")
    
    # Step 2: Copy ONLY tables from all files
    print("\n[2/3] Extracting tables from all 16 files...")
    total_tables = 0
    total_images = 0
    
    # Track image relationships
    image_rels = {}
    
    for week in range(1, 17):
        filename = f"logbook minggu {week}.docx"
        
        if not os.path.exists(filename):
            print(f"   Week {week:2d}: ⚠ File not found, skipping...")
            continue
        
        doc = Document(filename)
        week_tables = 0
        week_images = 0
        
        # Build image relationship mapping
        rel_mapping = {}
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    image_hash = hash(image_bytes)
                    
                    if image_hash not in image_rels:
                        # Add image to merged doc
                        new_rel = merged_doc.part.relate_to(image_part, rel.reltype)
                        image_rels[image_hash] = new_rel.rId
                        week_images += 1
                    
                    rel_mapping[rel_id] = image_rels[image_hash]
                except:
                    pass
        
        # Copy ONLY table elements (skip all paragraphs)
        for element in doc.element.body:
            if element.tag.endswith('tbl'):
                # Add small spacing before table
                merged_doc.add_paragraph()
                
                # Deep copy table
                new_table = deepcopy(element)
                
                # Update image relationship IDs
                from docx.oxml.ns import qn
                for blip in new_table.findall('.//' + qn('a:blip')):
                    old_embed = blip.get(qn('r:embed'))
                    if old_embed and old_embed in rel_mapping:
                        blip.set(qn('r:embed'), rel_mapping[old_embed])
                
                merged_doc.element.body.append(new_table)
                week_tables += 1
                total_tables += 1
        
        total_images += week_images
        print(f"   Week {week:2d}: {week_tables} table(s), {week_images} image(s)")
    
    print(f"\n   Total: {total_tables} tables, {total_images} images")
    
    # Step 3: Save to NEW file
    print("\n[3/3] Saving to new file...")
    output_file = "Logbook Lengkap - Clean.docx"
    merged_doc.save(output_file)
    
    # Verify
    verify_doc = Document(output_file)
    verify_tables = len(verify_doc.tables)
    
    # Count header paragraphs
    header_keywords = ["Nama", "NIM", "Program Studi"]
    header_paras = sum(1 for p in verify_doc.paragraphs if any(k in p.text for k in header_keywords))
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"New file created: {output_file}")
    print(f"  - Header elements: {header_count}")
    print(f"  - Header paragraphs with keywords: {header_paras}")
    print(f"  - Total tables: {verify_tables}")
    print(f"  - Total images: {total_images}")
    print("="*70)

if __name__ == "__main__":
    create_clean_merged_logbook()
