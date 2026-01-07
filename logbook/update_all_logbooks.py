from docx import Document
import os

# Define 16-week timeline with activities and results
TIMELINE = {
    1: [
        {
            "kegiatan": "Mengumpulkan kebutuhan sistem untuk aplikasi pemesanan warung sate berbasis web.\n\nMengidentifikasi fitur utama: admin panel, customer ordering, QR code per meja.",
            "hasil": "Memahami ruang lingkup proyek secara menyeluruh.\n\nMenemukan komponen inti yang wajib ada sebelum masuk tahap desain teknis."
        },
        {
            "kegiatan": "Menyusun Use Case Diagram untuk memetakan peran admin dan customer dalam sistem.",
            "hasil": "Mendapat gambaran jelas mengenai interaksi antara user dengan sistem dan fungsi yang harus disediakan."
        }
    ],
    2: [
        {
            "kegiatan": "Merancang struktur database untuk tabel users, menus, tables, orders, dan payment_methods.",
            "hasil": "Skema database yang terstruktur dan siap diimplementasikan."
        },
        {
            "kegiatan": "Setup project Laravel dengan konfigurasi awal dan instalasi dependencies.",
            "hasil": "Project Laravel berhasil diinisialisasi dan siap untuk development."
        }
    ],
    3: [
        {
            "kegiatan": "Membuat migration untuk tabel users, cache, dan jobs sebagai fondasi sistem.",
            "hasil": "Database schema dasar berhasil dibuat dan dapat dijalankan."
        },
        {
            "kegiatan": "Implementasi sistem authentication menggunakan Laravel Breeze.",
            "hasil": "Fitur login dan logout untuk admin berhasil diimplementasikan."
        }
    ],
    4: [
        {
            "kegiatan": "Membuat migration untuk tabel menus dengan field name, description, price, category, dan image.",
            "hasil": "Struktur tabel menu berhasil dibuat di database."
        },
        {
            "kegiatan": "Membuat Model Menu dengan relationship dan validasi data.",
            "hasil": "Model Menu siap digunakan untuk operasi CRUD."
        }
    ],
    5: [
        {
            "kegiatan": "Implementasi CRUD menu di admin panel dengan form create, edit, dan delete.",
            "hasil": "Admin dapat mengelola data menu dengan lengkap."
        },
        {
            "kegiatan": "Membuat tampilan daftar menu dengan pagination dan search functionality.",
            "hasil": "Daftar menu dapat ditampilkan dan dicari dengan mudah."
        }
    ],
    6: [
        {
            "kegiatan": "Upload dan validasi gambar menu menggunakan storage Laravel.",
            "hasil": "Gambar menu dapat diupload dan ditampilkan dengan benar."
        },
        {
            "kegiatan": "Implementasi kategori menu (makanan dan minuman) dengan filter.",
            "hasil": "Menu dapat dikelompokkan berdasarkan kategori."
        }
    ],
    7: [
        {
            "kegiatan": "Membuat migration untuk tabel tables dengan field table_number, uuid, status, dan location.",
            "hasil": "Struktur tabel untuk manajemen meja berhasil dibuat."
        },
        {
            "kegiatan": "Implementasi CRUD table management di admin panel.",
            "hasil": "Admin dapat mengelola data meja dengan status available/occupied."
        }
    ],
    8: [
        {
            "kegiatan": "Generate UUID unik untuk setiap meja sebagai identifier QR code.",
            "hasil": "Setiap meja memiliki UUID yang dapat digunakan untuk QR code."
        },
        {
            "kegiatan": "Implementasi fitur location validation untuk memastikan customer berada di lokasi warung.",
            "hasil": "Sistem dapat memvalidasi lokasi customer saat melakukan pemesanan."
        }
    ],
    9: [
        {
            "kegiatan": "Membuat halaman customer untuk menampilkan menu berdasarkan UUID meja.",
            "hasil": "Customer dapat melihat menu setelah scan QR code meja."
        },
        {
            "kegiatan": "Implementasi shopping cart functionality untuk menambah item pesanan.",
            "hasil": "Customer dapat menambahkan menu ke keranjang belanja."
        }
    ],
    10: [
        {
            "kegiatan": "Membuat session management untuk menyimpan data cart customer.",
            "hasil": "Data keranjang belanja tersimpan selama session aktif."
        },
        {
            "kegiatan": "Implementasi update quantity dan remove item dari cart.",
            "hasil": "Customer dapat mengubah jumlah atau menghapus item dari keranjang."
        }
    ],
    11: [
        {
            "kegiatan": "Membuat migration untuk tabel orders dengan field table_id, total_price, status, dan payment_method.",
            "hasil": "Struktur tabel orders berhasil dibuat di database."
        },
        {
            "kegiatan": "Membuat migration untuk tabel order_items sebagai detail pesanan.",
            "hasil": "Relasi antara orders dan menu items berhasil dibuat."
        }
    ],
    12: [
        {
            "kegiatan": "Implementasi checkout flow dengan kalkulasi total harga otomatis.",
            "hasil": "Customer dapat melakukan checkout dan melihat total pembayaran."
        },
        {
            "kegiatan": "Membuat migration untuk tabel payment_methods dan integrasi dengan order.",
            "hasil": "Sistem pembayaran dapat mendukung multiple payment methods."
        }
    ],
    13: [
        {
            "kegiatan": "Membuat admin dashboard untuk melihat daftar semua pesanan.",
            "hasil": "Admin dapat memantau semua pesanan yang masuk secara real-time."
        },
        {
            "kegiatan": "Implementasi update status pesanan (pending, processing, completed, cancelled).",
            "hasil": "Admin dapat mengubah status pesanan sesuai progress."
        }
    ],
    14: [
        {
            "kegiatan": "Implementasi payment verification untuk konfirmasi pembayaran customer.",
            "hasil": "Admin dapat memverifikasi pembayaran yang dilakukan customer."
        },
        {
            "kegiatan": "Membuat halaman order status untuk customer melihat progress pesanan.",
            "hasil": "Customer dapat melihat status pesanan mereka secara real-time."
        }
    ],
    15: [
        {
            "kegiatan": "Testing end-to-end flow dari customer order hingga admin verification.",
            "hasil": "Semua fitur utama berjalan dengan baik tanpa error kritis."
        },
        {
            "kegiatan": "Bug fixing untuk validasi form, error handling, dan edge cases.",
            "hasil": "Sistem lebih stabil dan user-friendly."
        }
    ],
    16: [
        {
            "kegiatan": "Membuat dokumentasi user manual untuk admin dan customer.",
            "hasil": "Dokumentasi lengkap untuk penggunaan sistem."
        },
        {
            "kegiatan": "Persiapan deployment dan testing di environment production.",
            "hasil": "Aplikasi siap untuk di-deploy dan digunakan."
        }
    ]
}

def update_logbook(week_number):
    """Update a single logbook file"""
    filename = f"logbook minggu {week_number}.docx"
    
    if not os.path.exists(filename):
        print(f"❌ File {filename} tidak ditemukan!")
        return False
    
    try:
        # Open document
        doc = Document(filename)
        
        # Get the main table
        if len(doc.tables) == 0:
            print(f"❌ Tidak ada tabel di {filename}")
            return False
        
        table = doc.tables[0]
        
        # Get activities for this week
        activities = TIMELINE.get(week_number, [])
        
        # Update rows (skip header row 0)
        for idx, activity in enumerate(activities):
            row_idx = idx + 1  # Start from row 1 (row 0 is header)
            
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                
                # Update Column 2 (Kegiatan)
                if len(row.cells) > 2:
                    row.cells[2].text = activity["kegiatan"]
                
                # Update Column 3 (Hal yang Diperoleh)
                if len(row.cells) > 3:
                    row.cells[3].text = activity["hasil"]
        
        # Save the document (overwrite)
        doc.save(filename)
        print(f"✅ Berhasil update {filename}")
        return True
        
    except Exception as e:
        print(f"❌ Error saat update {filename}: {str(e)}")
        return False

def main():
    print("="*70)
    print("MEMULAI UPDATE LOGBOOK FILES")
    print("="*70)
    print()
    
    success_count = 0
    failed_count = 0
    
    for week in range(1, 17):
        if update_logbook(week):
            success_count += 1
        else:
            failed_count += 1
    
    print()
    print("="*70)
    print("SUMMARY")
    print("="*70)
    print(f"✅ Berhasil: {success_count} files")
    print(f"❌ Gagal: {failed_count} files")
    print()

if __name__ == "__main__":
    main()
