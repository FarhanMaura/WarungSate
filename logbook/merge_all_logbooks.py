from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import os

def merge_logbooks_final():
    """Merge all 16 logbook files with proper image handling"""
    print("="*70)
    print("MERGING LOGBOOK FILES (Final Version)")
    print("="*70)
    
    # Create new document
    merged_doc = Document()
    
    # Track image parts to avoid duplicates
    image_parts = {}
    
    # Step 1: Copy header from first file
    print("\n[1/4] Copying header section from logbook minggu 1...")
    first_file = "logbook minggu 1.docx"
    
    if not os.path.exists(first_file):
        print(f"ERROR: {first_file} not found!")
        return
    
    doc1 = Document(first_file)
    
    # Copy header paragraphs (everything before first table)
    header_count = 0
    
    for element in doc1.element.body:
        if element.tag.endswith('tbl'):
            # Found first table, stop copying header
            break
        elif element.tag.endswith('p'):
            # Copy paragraph element using deep copy
            new_element = deepcopy(element)
            merged_doc.element.body.append(new_element)
            header_count += 1
    
    print(f"   ✓ Copied {header_count} header elements")
    
    # Step 2: Process each week's file
    print("\n[2/4] Copying tables and images from all 16 files...")
    total_tables = 0
    total_images = 0
    
    for week in range(1, 17):
        filename = f"logbook minggu {week}.docx"
        
        if not os.path.exists(filename):
            print(f"   ⚠ WARNING: {filename} not found, skipping...")
            continue
        
        doc = Document(filename)
        table_count = 0
        week_images = 0
        
        # Build relationship ID mapping for this document
        rel_mapping = {}
        
        # First, copy all image parts from this document
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    # Check if we already have this image
                    image_hash = hash(image_bytes)
                    
                    if image_hash not in image_parts:
                        # Add new image to merged document
                        new_image_part = merged_doc.part.package.image_parts._get_by_sha1(
                            image_part.sha1
                        )
                        
                        if new_image_part is None:
                            # Image doesn't exist, add it
                            new_rel = merged_doc.part.relate_to(image_part, rel.reltype)
                            rel_mapping[rel_id] = new_rel.rId
                            image_parts[image_hash] = new_rel.rId
                            week_images += 1
                        else:
                            # Image exists, reuse it
                            for r_id, r in merged_doc.part.rels.items():
                                if r.target_part == new_image_part:
                                    rel_mapping[rel_id] = r_id
                                    break
                    else:
                        # Reuse existing image
                        rel_mapping[rel_id] = image_parts[image_hash]
                        
                except Exception as e:
                    print(f"      Note: Error processing image: {e}")
        
        # Now copy all table elements from this document
        for element in doc.element.body:
            if element.tag.endswith('tbl'):
                # Add spacing paragraph before table
                merged_doc.add_paragraph()
                
                # Deep copy the entire table element
                new_table_element = deepcopy(element)
                
                # Update image relationship IDs in the copied table
                for blip in new_table_element.findall('.//' + qn('a:blip')):
                    old_embed = blip.get(qn('r:embed'))
                    if old_embed and old_embed in rel_mapping:
                        blip.set(qn('r:embed'), rel_mapping[old_embed])
                
                # Append to merged document
                merged_doc.element.body.append(new_table_element)
                
                table_count += 1
                total_tables += 1
        
        total_images += week_images
        print(f"   Week {week:2d}: {table_count} table(s), {week_images} image(s) from {filename}")
    
    print(f"\n   Total tables copied: {total_tables}")
    print(f"   Total unique images: {total_images}")
    
    # Step 3: Save merged document
    print("\n[3/4] Saving merged document...")
    output_file = "Logbook Lengkap.docx"
    merged_doc.save(output_file)
    
    # Step 4: Verify
    print("\n[4/4] Verifying merged document...")
    verify_doc = Document(output_file)
    verify_tables = len(verify_doc.tables)
    verify_images = 0
    
    try:
        for rel in verify_doc.part.rels.values():
            if "image" in rel.target_ref:
                verify_images += 1
    except:
        pass
    
    print(f"   Tables in merged file: {verify_tables}")
    print(f"   Images in merged file: {verify_images}")
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"Merged logbook saved as: {output_file}")
    print(f"  - Header elements: {header_count}")
    print(f"  - Total tables: {total_tables}")
    print(f"  - Total images: {total_images}")
    print(f"  - Original 16 files: PRESERVED")
    print("="*70)

if __name__ == "__main__":
    merge_logbooks_final()
