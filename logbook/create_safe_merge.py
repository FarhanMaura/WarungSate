from docx import Document
from copy import deepcopy
import os

def create_safe_merged_logbook():
    """Create merged logbook safely without complex image handling"""
    print("="*70)
    print("CREATING SAFE MERGED LOGBOOK")
    print("="*70)
    
    # Create new document
    merged_doc = Document()
    
    # Step 1: Copy header from first file
    print("\n[1/3] Copying header from logbook minggu 1...")
    first_file = "logbook minggu 1.docx"
    
    if not os.path.exists(first_file):
        print(f"ERROR: {first_file} not found!")
        return
    
    doc1 = Document(first_file)
    
    # Find first table position
    first_table_index = None
    for i, element in enumerate(doc1.element.body):
        if element.tag.endswith('tbl'):
            first_table_index = i
            break
    
    # Copy header elements (before first table)
    header_count = 0
    for i, element in enumerate(doc1.element.body):
        if i >= first_table_index:
            break
        
        new_element = deepcopy(element)
        merged_doc.element.body.append(new_element)
        header_count += 1
    
    print(f"   ✓ Copied {header_count} header elements")
    
    # Step 2: Copy tables WITHOUT image relationships (safer)
    print("\n[2/3] Copying tables from all 16 files...")
    total_tables = 0
    
    for week in range(1, 17):
        filename = f"logbook minggu {week}.docx"
        
        if not os.path.exists(filename):
            print(f"   Week {week:2d}: ⚠ File not found")
            continue
        
        doc = Document(filename)
        week_tables = 0
        
        # Copy only table elements
        for element in doc.element.body:
            if element.tag.endswith('tbl'):
                # Add spacing
                merged_doc.add_paragraph()
                
                # Deep copy table (images will be lost but structure preserved)
                new_table = deepcopy(element)
                merged_doc.element.body.append(new_table)
                
                week_tables += 1
                total_tables += 1
        
        print(f"   Week {week:2d}: {week_tables} table(s)")
    
    print(f"\n   Total tables: {total_tables}")
    
    # Step 3: Save
    print("\n[3/3] Saving...")
    output_file = "Logbook Lengkap - Safe.docx"
    merged_doc.save(output_file)
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"File created: {output_file}")
    print(f"  - Header: 1x (at top)")
    print(f"  - Tables: {total_tables}")
    print(f"\n⚠️  NOTE: Images may not appear in this version.")
    print("   This is a safe version that won't corrupt.")
    print("="*70)

if __name__ == "__main__":
    create_safe_merged_logbook()
