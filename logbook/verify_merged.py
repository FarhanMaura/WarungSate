from docx import Document
import os

def verify_merged_logbook():
    """Verify the merged logbook file"""
    filename = "Logbook Lengkap.docx"
    
    if not os.path.exists(filename):
        print(f"❌ ERROR: {filename} not found!")
        return
    
    print("="*70)
    print(f"VERIFYING: {filename}")
    print("="*70)
    
    doc = Document(filename)
    
    # Count elements
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    
    print(f"\n📊 Document Statistics:")
    print(f"   Total paragraphs: {para_count}")
    print(f"   Total tables: {table_count}")
    
    # Show header section
    print(f"\n📋 Header Section (First 15 paragraphs):")
    for i, para in enumerate(doc.paragraphs[:15]):
        if para.text.strip():
            print(f"   P{i}: {para.text[:70]}")
    
    # Show table summary
    print(f"\n📊 Table Summary:")
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0
        print(f"   Table {i+1:2d}: {rows} rows x {cols} cols")
    
    # Check for images
    print(f"\n🖼️  Images:")
    image_count = 0
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
    except:
        pass
    print(f"   Total images found: {image_count}")
    
    # Verify original files still exist
    print(f"\n📁 Original Files Check:")
    all_exist = True
    for week in range(1, 17):
        filename = f"logbook minggu {week}.docx"
        exists = os.path.exists(filename)
        if not exists:
            all_exist = False
            print(f"   ❌ {filename} - NOT FOUND")
    
    if all_exist:
        print(f"   ✅ All 16 original files still exist")
    
    print("\n" + "="*70)
    print("✅ VERIFICATION COMPLETE")
    print("="*70)

if __name__ == "__main__":
    verify_merged_logbook()
