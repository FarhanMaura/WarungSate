"""
Final approach: Use docxcompose to merge (preserves images properly),
then manually remove duplicate headers in a separate step.
"""
from docx import Document
from docxcompose.composer import Composer
import os

def final_merge_solution():
    """Merge with docxcompose, then clean headers"""
    print("="*70)
    print("FINAL MERGE SOLUTION")
    print("="*70)
    
    # Step 1: Merge all files with docxcompose
    print("\n[1/2] Merging all 16 files with docxcompose...")
    
    first_file = "logbook minggu 1.docx"
    if not os.path.exists(first_file):
        print(f"ERROR: {first_file} not found!")
        return
    
    master = Document(first_file)
    composer = Composer(master)
    
    for week in range(2, 17):
        filename = f"logbook minggu {week}.docx"
        if os.path.exists(filename):
            doc = Document(filename)
            composer.append(doc)
            print(f"   ✓ Week {week}")
    
    temp_file = "temp_full_merge.docx"
    composer.save(temp_file)
    print(f"   ✓ Saved temporary merge")
    
    # Step 2: Remove duplicate headers aggressively
    print("\n[2/2] Removing duplicate headers...")
    
    doc = Document(temp_file)
    
    # Find first table
    first_table_pos = None
    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('tbl'):
            first_table_pos = i
            break
    
    # Header indicators
    indicators = [
        "Nama", "NIM", "Program Studi", "Nomor HP",
        "Dosen Pembimbing", "Lokasi Pelaksanaan", "Waktu Pelaksanaan",
        "PROJECT INDEPENDENT", "Muhamad Farhan", "221420075",
        "Teknik Informatika", "083826383761", "Rasmila",
        "Warung Sate Madura", "06 Oktober 2025", "19 Januari 2026"
    ]
    
    # Remove paragraphs after first table that contain indicators
    elements_to_remove = []
    
    for i in range(first_table_pos + 1, len(doc.element.body)):
        element = doc.element.body[i]
        
        if element.tag.endswith('p'):
            para_index = None
            for j, p in enumerate(doc.paragraphs):
                if p._element == element:
                    para_index = j
                    break
            
            if para_index is not None:
                text = doc.paragraphs[para_index].text.strip()
                if any(ind in text for ind in indicators):
                    elements_to_remove.append(element)
    
    print(f"   Removing {len(elements_to_remove)} duplicate header paragraphs...")
    
    for element in elements_to_remove:
        element.getparent().remove(element)
    
    # Save final file
    output_file = "Logbook Lengkap Final.docx"
    doc.save(output_file)
    
    # Clean up temp
    if os.path.exists(temp_file):
        os.remove(temp_file)
    
    # Verify
    verify_doc = Document(output_file)
    tables = len(verify_doc.tables)
    
    images = 0
    try:
        for rel in verify_doc.part.rels.values():
            if "image" in rel.target_ref:
                images += 1
    except:
        pass
    
    # Count remaining headers
    remaining_headers = 0
    for para in verify_doc.paragraphs:
        if any(ind in para.text for ind in indicators):
            remaining_headers += 1
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"File created: {output_file}")
    print(f"  - Tables: {tables}")
    print(f"  - Images: {images}")
    print(f"  - Header paragraphs remaining: {remaining_headers}")
    print("="*70)
    
    print(f"\n💡 Try opening: {output_file}")
    print("   If it asks to recover, click 'Yes'")

if __name__ == "__main__":
    final_merge_solution()
