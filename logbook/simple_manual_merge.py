from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os
import io

def copy_paragraph_to_cell(source_para, target_cell):
    """Copy paragraph content to target cell"""
    target_para = target_cell.add_paragraph()
    target_para.alignment = source_para.alignment
    
    for run in source_para.runs:
        new_run = target_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name
        
        # Try to copy images
        for drawing in run._element.findall('.//' + qn('w:drawing')):
            try:
                blip = drawing.find('.//' + qn('a:blip'))
                if blip is not None:
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        # Get image from source
                        image_part = source_para._element.part.related_parts[embed]
                        image_bytes = image_part.blob
                        
                        # Add to target
                        new_run.add_picture(io.BytesIO(image_bytes))
            except:
                pass

def simple_manual_merge():
    """Simple manual merge - copy everything manually"""
    print("="*70)
    print("SIMPLE MANUAL MERGE")
    print("="*70)
    
    merged_doc = Document()
    
    # Step 1: Copy header from first file
    print("\n[1/3] Copying header...")
    first_file = "logbook minggu 1.docx"
    
    if not os.path.exists(first_file):
        print(f"ERROR: {first_file} not found!")
        return
    
    doc1 = Document(first_file)
    
    # Copy paragraphs until first table
    for para in doc1.paragraphs:
        # Check if we've reached a table
        if para._element.getnext() is not None and para._element.getnext().tag.endswith('tbl'):
            # Copy this paragraph and stop
            new_para = merged_doc.add_paragraph(para.text)
            new_para.alignment = para.alignment
            break
        
        # Copy paragraph
        new_para = merged_doc.add_paragraph(para.text)
        new_para.alignment = para.alignment
        
        # Copy formatting
        if para.runs:
            new_para.clear()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
    
    print("   ✓ Header copied")
    
    # Step 2: Copy tables from all files
    print("\n[2/3] Copying tables...")
    total_tables = 0
    
    for week in range(1, 17):
        filename = f"logbook minggu {week}.docx"
        
        if not os.path.exists(filename):
            print(f"   Week {week:2d}: File not found")
            continue
        
        doc = Document(filename)
        week_tables = 0
        
        for table in doc.tables:
            # Add spacing
            merged_doc.add_paragraph()
            
            # Create new table with same dimensions
            rows = len(table.rows)
            cols = len(table.rows[0].cells) if rows > 0 else 0
            
            new_table = merged_doc.add_table(rows=rows, cols=cols)
            
            # Copy table style
            try:
                if table.style:
                    new_table.style = table.style
            except:
                pass
            
            # Copy each cell
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    target_cell = new_table.rows[i].cells[j]
                    
                    # Clear default paragraph
                    if target_cell.paragraphs:
                        target_cell.paragraphs[0].clear()
                    
                    # Copy all paragraphs
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para_idx == 0 and target_cell.paragraphs:
                            target_para = target_cell.paragraphs[0]
                        else:
                            target_para = target_cell.add_paragraph()
                        
                        target_para.alignment = para.alignment
                        
                        # Copy runs
                        for run in para.runs:
                            new_run = target_para.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.name:
                                new_run.font.name = run.font.name
                            
                            # Try copy images
                            try:
                                for drawing in run._element.findall('.//' + qn('w:drawing')):
                                    blip = drawing.find('.//' + qn('a:blip'))
                                    if blip is not None:
                                        embed = blip.get(qn('r:embed'))
                                        if embed:
                                            image_part = cell._tc.part.related_parts[embed]
                                            image_bytes = image_part.blob
                                            new_run.add_picture(io.BytesIO(image_bytes))
                            except:
                                pass
            
            week_tables += 1
            total_tables += 1
        
        print(f"   Week {week:2d}: {week_tables} table(s)")
    
    # Step 3: Save
    print("\n[3/3] Saving...")
    output_file = "Logbook Lengkap Manual.docx"
    merged_doc.save(output_file)
    
    print(f"\n{'='*70}")
    print("✅ DONE!")
    print(f"{'='*70}")
    print(f"File: {output_file}")
    print(f"Tables: {total_tables}")
    print("="*70)

if __name__ == "__main__":
    simple_manual_merge()
