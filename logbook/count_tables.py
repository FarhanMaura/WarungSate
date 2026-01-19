from docx import Document
import os

def count_tables_per_week():
    """Count tables in each weekly logbook file"""
    print("="*70)
    print("COUNTING TABLES PER WEEK")
    print("="*70)
    
    total_tables = 0
    
    for week in range(1, 17):
        filename = f"logbook minggu {week}.docx"
        
        if not os.path.exists(filename):
            print(f"Week {week:2d}: FILE NOT FOUND")
            continue
        
        doc = Document(filename)
        table_count = len(doc.tables)
        total_tables += table_count
        
        print(f"Week {week:2d}: {table_count} table(s)")
    
    print("="*70)
    print(f"TOTAL EXPECTED TABLES: {total_tables}")
    print("="*70)
    
    # Check merged file
    merged_file = "Logbook Lengkap.docx"
    if os.path.exists(merged_file):
        merged_doc = Document(merged_file)
        merged_tables = len(merged_doc.tables)
        print(f"\nMerged file has: {merged_tables} tables")
        
        if merged_tables < total_tables:
            print(f"⚠️  MISSING {total_tables - merged_tables} tables!")
        else:
            print("✅ All tables present!")

if __name__ == "__main__":
    count_tables_per_week()
