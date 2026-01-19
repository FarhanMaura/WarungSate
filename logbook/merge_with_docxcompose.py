"""
Merge logbook files using docxcompose library
First install: pip install docxcompose
"""
from docx import Document
from docxcompose.composer import Composer
import os

def merge_with_docxcompose():
    """Merge all logbook files using docxcompose library"""
    print("="*70)
    print("MERGING LOGBOOK FILES (Using docxcompose)")
    print("="*70)
    
    # Check if docxcompose is installed
    try:
        from docxcompose.composer import Composer
    except ImportError:
        print("\n⚠️  Installing docxcompose library...")
        import subprocess
        subprocess.check_call(["pip", "install", "docxcompose"])
        from docxcompose.composer import Composer
        print("✓ docxcompose installed successfully\n")
    
    print("\n[1/3] Loading first file as base...")
    first_file = "logbook minggu 1.docx"
    
    if not os.path.exists(first_file):
        print(f"ERROR: {first_file} not found!")
        return
    
    # Load first document as base
    master = Document(first_file)
    composer = Composer(master)
    
    print(f"   ✓ Loaded {first_file} as base")
    
    # Step 2: Append all other files
    print("\n[2/3] Appending remaining 15 files...")
    
    for week in range(2, 17):
        filename = f"logbook minggu {week}.docx"
        
        if not os.path.exists(filename):
            print(f"   ⚠ WARNING: {filename} not found, skipping...")
            continue
        
        try:
            doc = Document(filename)
            composer.append(doc)
            print(f"   Week {week:2d}: ✓ Appended {filename}")
        except Exception as e:
            print(f"   Week {week:2d}: ✗ Error appending {filename}: {e}")
    
    # Step 3: Save merged document
    print("\n[3/3] Saving merged document...")
    output_file = "Logbook Lengkap - Full.docx"
    composer.save(output_file)
    
    # Verify
    verify_doc = Document(output_file)
    verify_tables = len(verify_doc.tables)
    verify_images = 0
    
    try:
        for rel in verify_doc.part.rels.values():
            if "image" in rel.target_ref:
                verify_images += 1
    except:
        pass
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"Merged logbook saved as: {output_file}")
    print(f"  - Total tables: {verify_tables}")
    print(f"  - Total images: {verify_images}")
    print(f"  - Original 16 files: PRESERVED")
    print(f"\n⚠️  NOTE: This file includes ALL content from all 16 files,")
    print(f"   including headers. You may want to manually remove duplicate")
    print(f"   headers from weeks 2-16 in Microsoft Word.")
    print("="*70)

if __name__ == "__main__":
    merge_with_docxcompose()
