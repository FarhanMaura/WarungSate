from docx import Document

def analyze_document(filename):
    print(f"Analyzing {filename}...")
    try:
        doc = Document(filename)
        
        print("\n--- Searching for Header Fields in Paragraphs ---")
        targets = ["Nama", "NIM", "Program Studi", "Nomor HP", "Dosen Pembimbing", "Lokasi Pelaksanaan", "Waktu Pelaksanaan"]
        found = False
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if any(t in text for t in targets):
                print(f"P{i}: {text}")
                found = True
            elif i < 15 and text: # Print first 15 non-empty paragraphs anyway just to see
                print(f"P{i} (Context): {text}")
        
        if not found:
            print("Header fields not found in paragraphs. Checking tables...")
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if any(t in row_text for t in targets):
                        print(f"Table {t_idx} Row {r_idx}: {row_text}")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    analyze_document("logbook minggu 1.docx")
