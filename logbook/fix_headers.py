from docx import Document
import os

def fix_duplicate_headers():
    """Remove ALL duplicate headers, keeping only the first one"""
    print("="*70)
    print("FIXING DUPLICATE HEADERS")
    print("="*70)
    
    input_file = "Logbook Lengkap.docx"
    
    if not os.path.exists(input_file):
        print(f"ERROR: {input_file} not found!")
        return
    
    print(f"\n[1/3] Loading {input_file}...")
    doc = Document(input_file)
    
    # Keywords that identify header sections
    header_keywords = [
        "Nama", "NIM", "Program Studi", "Nomor HP", 
        "Dosen Pembimbing", "Lokasi Pelaksanaan", 
        "Waktu Pelaksanaan", "PROJECT INDEPENDENT"
    ]
    
    print("\n[2/3] Identifying and removing duplicate headers...")
    
    elements_to_remove = []
    first_table_index = None
    
    # Find index of first table
    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('tbl'):
            first_table_index = i
            print(f"   First table found at index {i}")
            break
    
    if first_table_index is None:
        print("   ERROR: No tables found!")
        return
    
    # Remove ALL paragraphs after first table that contain header keywords
    removed_count = 0
    
    for i, element in enumerate(doc.element.body):
        if i > first_table_index and element.tag.endswith('p'):
            # Get paragraph text
            para_index = [j for j, p in enumerate(doc.paragraphs) if p._element == element]
            if para_index:
                para = doc.paragraphs[para_index[0]]
                text = para.text.strip()
                
                # Check if this paragraph contains any header keyword
                if any(keyword in text for keyword in header_keywords):
                    elements_to_remove.append(element)
                    removed_count += 1
                    if removed_count <= 20:  # Show first 20 removals
                        print(f"   Removing: {text[:60]}")
    
    print(f"\n   Total paragraphs to remove: {removed_count}")
    
    # Remove marked elements
    for element in elements_to_remove:
        element.getparent().remove(element)
    
    # Save cleaned document
    print(f"\n[3/3] Saving cleaned document...")
    doc.save(input_file)
    
    # Verify
    verify_doc = Document(input_file)
    verify_tables = len(verify_doc.tables)
    
    # Count remaining header paragraphs
    remaining_headers = 0
    for para in verify_doc.paragraphs:
        if any(kw in para.text for kw in header_keywords):
            remaining_headers += 1
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"  - Removed: {removed_count} duplicate header paragraphs")
    print(f"  - Remaining header paragraphs: {remaining_headers}")
    print(f"  - Total tables: {verify_tables}")
    print("="*70)
    
    if remaining_headers > 10:
        print(f"\n⚠️  WARNING: Still {remaining_headers} header paragraphs remaining.")
        print("   Expected around 7-8 (one set of headers).")

if __name__ == "__main__":
    fix_duplicate_headers()
