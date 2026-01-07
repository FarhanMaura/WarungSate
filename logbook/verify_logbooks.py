from docx import Document

def verify_logbook(filename, week_number):
    """Verify a logbook file has been updated correctly"""
    print(f"\n{'='*70}")
    print(f"VERIFYING: {filename} (Minggu {week_number})")
    print(f"{'='*70}\n")
    
    doc = Document(filename)
    
    if len(doc.tables) == 0:
        print("❌ Tidak ada tabel!")
        return
    
    table = doc.tables[0]
    
    print(f"Total rows: {len(table.rows)}")
    print(f"Total columns: {len(table.columns)}\n")
    
    # Show header
    print("HEADER ROW:")
    for col_idx, cell in enumerate(table.rows[0].cells):
        print(f"  Column {col_idx}: {cell.text.strip()}")
    
    # Show data rows
    print("\nDATA ROWS:")
    for row_idx in range(1, min(len(table.rows), 4)):  # Show max 3 data rows
        row = table.rows[row_idx]
        print(f"\n--- Row {row_idx} ---")
        
        # Show all columns
        for col_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if col_idx == 2:  # Kegiatan
                print(f"  [KEGIATAN]: {text[:100]}...")
            elif col_idx == 3:  # Hasil
                print(f"  [HASIL]: {text[:100]}...")
            else:
                print(f"  Column {col_idx}: {text[:50]}")

if __name__ == "__main__":
    # Verify a few sample files
    print("\n" + "="*70)
    print("VERIFICATION REPORT")
    print("="*70)
    
    # Check week 1 (planning)
    verify_logbook("logbook minggu 1.docx", 1)
    
    # Check week 8 (table management)
    verify_logbook("logbook minggu 8.docx", 8)
    
    # Check week 16 (documentation)
    verify_logbook("logbook minggu 16.docx", 16)
    
    print("\n" + "="*70)
    print("VERIFICATION COMPLETE")
    print("="*70)
