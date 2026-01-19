from docx import Document
from docxcompose.composer import Composer
import os

def create_final_logbook():
    """Create final merged logbook with single header and all tables"""
    print("="*70)
    print("CREATING FINAL MERGED LOGBOOK")
    print("="*70)
    
    # Step 1: Merge all files first
    print("\n[1/4] Merging all 16 files...")
    first_file = "logbook minggu 1.docx"
    
    if not os.path.exists(first_file):
        print(f"ERROR: {first_file} not found!")
        return
    
    master = Document(first_file)
    composer = Composer(master)
    
    for week in range(2, 17):
        filename = f"logbook minggu {week}.docx"
        if os.path.exists(filename):
            doc = Document(filename)
            composer.append(doc)
            print(f"   ✓ Week {week}")
    
    temp_file = "temp_merged.docx"
    composer.save(temp_file)
    print(f"   ✓ Saved temporary merged file")
    
    # Step 2: Clean up duplicate headers
    print("\n[2/4] Removing duplicate headers...")
    doc = Document(temp_file)
    
    header_keywords = ["Nama", "NIM", "Program Studi", "PROJECT INDEPENDENT"]
    elements_to_remove = []
    first_header_passed = False
    consecutive_header_paras = 0
    
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Get paragraph text
            para_index = [j for j, p in enumerate(doc.paragraphs) if p._element == element]
            if para_index:
                para = doc.paragraphs[para_index[0]]
                text = para.text.strip()
                
                # Check if this looks like a header paragraph
                is_header = any(keyword in text for keyword in header_keywords)
                
                if is_header:
                    consecutive_header_paras += 1
                    if first_header_passed:
                        # This is a duplicate header
                        elements_to_remove.append(element)
                else:
                    # Reset counter if we see 5+ consecutive header-like paragraphs
                    if consecutive_header_paras >= 5:
                        first_header_passed = True
                    consecutive_header_paras = 0
        
        elif element.tag.endswith('tbl'):
            # Table found, definitely past first header
            first_header_passed = True
            consecutive_header_paras = 0
    
    print(f"   ✓ Removing {len(elements_to_remove)} duplicate header elements")
    
    for element in elements_to_remove:
        element.getparent().remove(element)
    
    # Step 3: Save final file
    print("\n[3/4] Saving final logbook...")
    output_file = "Logbook Lengkap.docx"
    doc.save(output_file)
    
    # Clean up temp file
    if os.path.exists(temp_file):
        os.remove(temp_file)
    
    # Step 4: Verify
    print("\n[4/4] Verifying final file...")
    verify_doc = Document(output_file)
    table_count = len(verify_doc.tables)
    image_count = 0
    
    try:
        for rel in verify_doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
    except:
        pass
    
    # Count paragraphs with header keywords
    header_para_count = 0
    for para in verify_doc.paragraphs:
        if any(kw in para.text for kw in header_keywords):
            header_para_count += 1
            if header_para_count <= 10:  # Show first 10 header-related paragraphs
                print(f"   Header: {para.text[:60]}")
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"Final logbook: {output_file}")
    print(f"  - Tables: {table_count}")
    print(f"  - Images: {image_count}")
    print(f"  - Header paragraphs: {header_para_count}")
    print(f"  - Original files: PRESERVED")
    print("="*70)
    
    if table_count < 30:
        print(f"\n⚠️  WARNING: Expected ~31 tables, found {table_count}")
        print("   Some weeks might have merged tables. Please verify manually.")
    
    return output_file

if __name__ == "__main__":
    create_final_logbook()
