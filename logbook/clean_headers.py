from docx import Document
import os

def remove_duplicate_headers():
    """Remove duplicate headers from merged logbook, keeping only the first one"""
    print("="*70)
    print("REMOVING DUPLICATE HEADERS")
    print("="*70)
    
    input_file = "Logbook Lengkap - Full.docx"
    output_file = "Logbook Lengkap.docx"
    
    if not os.path.exists(input_file):
        print(f"ERROR: {input_file} not found!")
        return
    
    print(f"\n[1/3] Loading {input_file}...")
    doc = Document(input_file)
    
    # Header keywords to identify header sections
    header_keywords = ["Nama", "NIM", "Program Studi", "Nomor HP", "Dosen Pembimbing", 
                      "Lokasi Pelaksanaan", "Waktu Pelaksanaan", "PROJECT INDEPENDENT"]
    
    # Find all header sections
    print("\n[2/3] Identifying and removing duplicate headers...")
    
    elements_to_remove = []
    in_header = False
    header_count = 0
    first_table_found = False
    
    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('p'):
            # Check if this paragraph is part of a header
            para_index = [j for j, p in enumerate(doc.paragraphs) if p._element == element]
            if para_index:
                para = doc.paragraphs[para_index[0]]
                text = para.text.strip()
                
                # Check if this is a header paragraph
                is_header_para = any(keyword in text for keyword in header_keywords)
                
                if is_header_para:
                    if first_table_found:
                        # This is a duplicate header (after first table), mark for removal
                        in_header = True
                        elements_to_remove.append(element)
                        header_count += 1
                    else:
                        # This is the first header, keep it
                        in_header = True
                elif in_header and not text:
                    # Empty paragraph in header section
                    if first_table_found:
                        elements_to_remove.append(element)
                elif in_header and text and not is_header_para:
                    # End of header section
                    in_header = False
        
        elif element.tag.endswith('tbl'):
            # Found a table
            first_table_found = True
            in_header = False
    
    # Remove marked elements
    print(f"   Found {header_count} duplicate header paragraphs")
    print(f"   Removing {len(elements_to_remove)} elements...")
    
    for element in elements_to_remove:
        element.getparent().remove(element)
    
    # Save cleaned document
    print(f"\n[3/3] Saving cleaned document as {output_file}...")
    doc.save(output_file)
    
    # Verify
    verify_doc = Document(output_file)
    verify_tables = len(verify_doc.tables)
    verify_images = 0
    
    try:
        for rel in verify_doc.part.rels.values():
            if "image" in rel.target_ref:
                verify_images += 1
    except:
        pass
    
    print(f"\n{'='*70}")
    print("✅ SUCCESS!")
    print(f"{'='*70}")
    print(f"Cleaned logbook saved as: {output_file}")
    print(f"  - Removed: {len(elements_to_remove)} duplicate header elements")
    print(f"  - Total tables: {verify_tables}")
    print(f"  - Total images: {verify_images}")
    print("="*70)

if __name__ == "__main__":
    remove_duplicate_headers()
