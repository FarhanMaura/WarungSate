from docx import Document
import sys

def extract_full_structure(filename):
    """Extract complete structure of logbook to understand what needs to be changed"""
    doc = Document(filename)
    
    print(f"\n{'='*70}")
    print(f"FULL STRUCTURE ANALYSIS: {filename}")
    print(f"{'='*70}\n")
    
    # Check tables
    print(f"Total Tables: {len(doc.tables)}\n")
    
    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f"Main Table: {len(table.rows)} rows x {len(table.columns)} columns\n")
        
        # Print header row
        print("="*70)
        print("TABLE STRUCTURE:")
        print("="*70)
        
        for row_idx, row in enumerate(table.rows):
            print(f"\n--- ROW {row_idx} ---")
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                # Show first 150 chars
                display_text = text[:150].replace('\n', ' | ')
                if text:
                    print(f"  Column {col_idx}: {display_text}")
                    if len(text) > 150:
                        print(f"    ... (total {len(text)} chars)")
        
        # Identify which columns contain "Kegiatan" and "Hasil"
        print("\n" + "="*70)
        print("COLUMN IDENTIFICATION:")
        print("="*70)
        
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for col_idx, cell in enumerate(header_row.cells):
                header_text = cell.text.strip().lower()
                print(f"Column {col_idx}: '{cell.text.strip()}'")
                if 'kegiatan' in header_text:
                    print(f"  -> This is KEGIATAN column")
                if 'hasil' in header_text:
                    print(f"  -> This is HASIL column")

if __name__ == "__main__":
    extract_full_structure("logbook minggu 1.docx")
