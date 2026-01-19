from docx import Document
import os

def verify_final_logbook():
    """Verify the final merged logbook"""
    print("="*70)
    print("FINAL VERIFICATION")
    print("="*70)
    
    filename = "Logbook Lengkap.docx"
    
    if not os.path.exists(filename):
        print(f"ERROR: {filename} not found!")
        return
    
    doc = Document(filename)
    
    # Header keywords
    header_keywords = [
        "Nama", "NIM", "Program Studi", "Nomor HP", 
        "Dosen Pembimbing", "Lokasi Pelaksanaan", 
        "Waktu Pelaksanaan"
    ]
    
    print(f"\n[1/3] Document Statistics:")
    print(f"   Total paragraphs: {len(doc.paragraphs)}")
    print(f"   Total tables: {len(doc.tables)}")
    
    # Count images
    image_count = 0
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
    except:
        pass
    print(f"   Total images: {image_count}")
    
    # Find header paragraphs
    print(f"\n[2/3] Header Analysis:")
    header_paras = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(kw in text for kw in header_keywords):
            header_paras.append((i, text))
    
    print(f"   Total header-related paragraphs: {len(header_paras)}")
    
    # Show first 10 header paragraphs
    print(f"\n   First 10 header paragraphs:")
    for i, (idx, text) in enumerate(header_paras[:10]):
        print(f"      P{idx}: {text[:60]}")
    
    # Check if there are duplicates after first table
    first_table_para_idx = None
    for i, element in enumerate(doc.element.body):
        if element.tag.endswith('tbl'):
            # Find paragraph index before this table
            para_count = 0
            for j, el in enumerate(doc.element.body[:i]):
                if el.tag.endswith('p'):
                    para_count += 1
            first_table_para_idx = para_count
            break
    
    if first_table_para_idx:
        headers_after_table = [p for p in header_paras if p[0] > first_table_para_idx]
        print(f"\n   Headers AFTER first table: {len(headers_after_table)}")
        
        if headers_after_table:
            print(f"   ⚠️  WARNING: Found {len(headers_after_table)} duplicate headers!")
            print(f"   First few duplicates:")
            for idx, text in headers_after_table[:5]:
                print(f"      P{idx}: {text[:60]}")
        else:
            print(f"   ✅ No duplicate headers found!")
    
    # Show table summary
    print(f"\n[3/3] Table Summary:")
    for i, table in enumerate(doc.tables[:5]):  # Show first 5 tables
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0
        print(f"   Table {i+1}: {rows} rows x {cols} cols")
    
    if len(doc.tables) > 5:
        print(f"   ... and {len(doc.tables) - 5} more tables")
    
    print(f"\n{'='*70}")
    print("VERIFICATION COMPLETE")
    print(f"{'='*70}")

if __name__ == "__main__":
    verify_final_logbook()
